import docx
import PyPDF2
import os

search_index = {}

def process_document(file_path):
    # Function to process a document based on its file type
    file_type = file_path.split('.')[-1]
    match file_type:
        case 'pdf':
            text = extract_text_from_pdf(file_path)
            index_document(file_path, text)
            return text
        case 'docx':
            text = extract_text_from_docx(file_path)
            index_document(file_path, text)
            return text
        case 'txt':
            text = extract_text_from_txt(file_path)
            index_document(file_path, text)
            return text
        case _:
            print(f"Unsupported file type: {file_type}")

def extract_text_from_pdf(pdf_path):
    # Function to extract text from a PDF file
    reader = PyPDF2.PdfReader(pdf_path)
    text = ""
    for page_num in range(len(reader.pages)):
        page = reader.pages[page_num]
        text += page.extract_text()
    return text


def extract_text_from_docx(docx_path):
    # Function to extract text from a DOCX file
    doc = docx.Document(docx_path)
    text = ""
    for para in doc.paragraphs:
        text += para.text
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text += cell.text
    return text

def extract_text_from_txt(txt_path):
    # Function to extract text from a TXT file
    with open(txt_path, 'r') as f:
        text = f.read()
    return text

def index_document(file_name, text):
    # Function to index the processed document
    search_index[file_name] = text
    print()
    print("Search index now contains:")
    print(f"Indexed document: {file_name}")
    print(f"Document text: {text}")
    print()

def handle_document_upload(file, file_path):
    # Copy the file to the docs folder
    fileType = file.filename.split('.')[-1]
    match fileType:
        case 'pdf':
            reader = PyPDF2.PdfReader(file)
            # save the pdf file to the file path using PyPDF2
            with open(file_path, 'wb') as f:
                writer = PyPDF2.PdfWriter()
                for page_num in range(len(reader.pages)):
                    writer.add_page(reader.pages[page_num])
                writer.write(f)
                writer.close()
        case 'docx':
            template = docx.Document()
            this_file = docx.Document(file)
            for para in this_file.paragraphs:
                template.add_paragraph(para.text)
            for table in this_file.tables:
                # Get the number of rows and columns from the table
                rows = len(table.rows)
                cols = len(table.columns)
                # Add a table to the template with the same number of rows and columns
                new_table = template.add_table(rows=rows, cols=cols)
                # Copy the content of each cell
                for i in range(rows):
                    for j in range(cols):
                        new_table.cell(i, j).text = table.cell(i, j).text
            template.save(file_path)
        case 'txt':
            fileContents = file.readlines()
            with open(file_path, 'w') as f:
                for line in fileContents:
                    f.write(line.decode('utf-8'))
        case _:
            print(f"Unsupported file type: {fileType}")
    # Process and index the document
    text = process_document(file_path)
    return text

def process_documents(files, target_directory):
    # Ensure the target directory exists
    target_directory = f"{os.path.dirname(__file__)}\\docs\\{target_directory}"
    os.makedirs(target_directory, exist_ok=True)

    for file in files:
        file_path = os.path.join(target_directory, file.filename)
        text = handle_document_upload(file, file_path)
    return search_index

def handle_documents(target_directory):
    # Process all documents in the target directory
    target_directory = f"{os.path.dirname(__file__)}/docs/{target_directory}"
    all_documents = []
    indexed_documents = {}

    for file_name in os.listdir(target_directory):
        if os.path.isfile(os.path.join(target_directory, file_name)):
            all_documents.append(file_name)

    for doc_name in all_documents:
        doc_path = os.path.join(target_directory, doc_name)
        text = process_document(doc_path)
        indexed_documents[doc_name] = text

    return indexed_documents