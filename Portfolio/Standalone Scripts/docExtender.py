from docx import Document

existing_text = []
# Example usage:
print("Please upload a Word Doc")

doc_filename = input("Enter path to word doc! ")

if doc_filename and doc_filename.lower().endswith('.docx'):
    doc = Document(doc_filename)
    print(f"Word Doc uploaded successfully: {doc_filename}")

    for paragraph in doc.paragraphs:
        existing_text.append(paragraph.text)

    #print("Existing Text:")
    #print(existing_text)

    newDoc = Document(doc_filename)
    for num in range(10):
      doc.add_paragraph(existing_text)
    additionalName = input("Enter additional name for new file: ")
    newName = additionalName + ".docx"
    doc_filename = doc_filename.replace(".docx", newName)
    doc.save(doc_filename)
else:
    print("Error: Please upload a Word Doc (.docx)")
