import docx
import PyPDF2
import os

search_index = {}

def process_document(file_path):
    # Function to process a document based on its file type
    file_type = file_path.split('.')[-1]
    match file_type:
        case 'pdf':
            text = extract_text_from_pdf(file_path)
            index_document(file_path, text)
            return text
        case 'docx':
            text = extract_text_from_docx(file_path)
            index_document(file_path, text)
            return text
        case 'txt':
            text = extract_text_from_txt(file_path)
            index_document(file_path, text)
            return text
        case _:
            print(f"Unsupported file type: {file_type}")

def extract_text_from_pdf(pdf_path):
    # Function to extract text from a PDF file
    reader = PyPDF2.PdfFileReader(pdf_path)
    text = ""
    for page_num in range(len(reader.pages)):
        page = reader.pages[page_num]
        text += page.extract_text()
    return text


def extract_text_from_docx(docx_path):
    # Function to extract text from a DOCX file
    doc = docx.Document(docx_path)
    text = ""
    for para in doc.paragraphs:
        text += para.text
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text += cell.text
    return text

def extract_text_from_txt(txt_path):
    # Function to extract text from a TXT file
    with open(txt_path, 'r') as f:
        text = f.read()
    return text

def index_document(file_name, text):
    # Function to index the processed document
    search_index[file_name] = text

def handle_document_upload(file):
    # Copy the file to the docs folder
    file_path = f"{os.path.dirname(__file__)}/docs/{file.filename}"
    fileType = file.filename.split('.')[-1]
    match fileType:
        case 'pdf':
            reader = PyPDF2.PdfReader(file)
            # save the pdf file to the file path using PyPDF2
            with open(file_path, 'wb') as f:
                writer = PyPDF2.PdfWriter()
                for page_num in range(len(reader.pages)):
                    writer.add_page(reader.pages[page_num])
                writer.write(f)
                writer.close()
        case 'docx':
            template = docx.Document()
            this_file = docx.Document(file)
            for para in this_file.paragraphs:
                template.add_paragraph(para.text)
            for table in this_file.tables:
                # Get the number of rows and columns from the table
                rows = len(table.rows)
                cols = len(table.columns)
                # Add a table to the template with the same number of rows and columns
                new_table = template.add_table(rows=rows, cols=cols)
                # Copy the content of each cell
                for i in range(rows):
                    for j in range(cols):
                        new_table.cell(i, j).text = table.cell(i, j).text
            template.save(file_path)
        case 'txt':
            fileContents = file.readlines()
            with open(file_path, 'w') as f:
                for line in fileContents:
                    f.write(line.decode('utf-8'))
        case _:
            print(f"Unsupported file type: {fileType}")
    # Process and index the document
    text = process_document(file_path)
    return text

def process_documents(files):
    # Function to process multiple documents
    document_text = []
    for file in files:
        text = handle_document_upload(file)
        document_text.append(text)
    return document_text, search_index